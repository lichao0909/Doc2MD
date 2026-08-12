"""生成一个示例 Word 文档用于演示转换。"""
from pathlib import Path

from docx import Document
from docx.shared import Pt


def build_docx(path: Path) -> None:
    doc = Document()

    doc.add_heading("项目开发计划书", level=0)

    doc.add_heading("1. 项目背景", level=1)
    doc.add_paragraph(
        "本项目旨在开发一个能够将 PDF 和 Word 文档批量转换为 Markdown 的命令行工具，"
        "方便在知识库、博客和技术文档中复用已有资料。"
    )

    doc.add_heading("2. 主要功能", level=1)
    features = [
        "支持 PDF、Word、Excel、PowerPoint、HTML 等格式",
        "保留标题层级、列表、表格结构",
        "支持批量转换与自定义输出目录",
        "命令行接口，可集成到脚本与 CI",
    ]
    for item in features:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. 里程碑", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "阶段"
    hdr[1].text = "时间"
    hdr[2].text = "交付物"
    rows = [
        ("需求确认", "第 1 周", "需求文档"),
        ("原型开发", "第 2-3 周", "可运行 CLI"),
        ("测试优化", "第 4 周", "测试报告"),
    ]
    for stage, time, deliverable in rows:
        cells = table.add_row().cells
        cells[0].text = stage
        cells[1].text = time
        cells[2].text = deliverable

    doc.add_heading("4. 备注", level=1)
    p = doc.add_paragraph()
    run = p.add_run("本文档由示例脚本自动生成，用于演示转换效果。")
    run.italic = True

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"已生成: {path}")


if __name__ == "__main__":
    build_docx(Path(__file__).parent / "demo.docx")
